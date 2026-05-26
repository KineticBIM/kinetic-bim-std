# -*- coding: utf-8 -*-
"""Generate the Auto Tag user guide as a Word .docx.

Run with the system Python that has python-docx installed (NOT the
IronPython runtime inside Revit). The output sits next to this script
and is intended for end users in the practice - BIM modellers and
coordinators - not for developers maintaining the package. Re-run any
time the tool changes; the generator is the source of truth, the
.docx is the artefact.

    > python build_user_guide.py
"""

from __future__ import print_function

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


HERE = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(HERE, "Auto Tag User Guide.docx")


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _set_font(run, name="Calibri", size=11, bold=False, italic=False,
              color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_mono(p, text):
    """Append a monospace span to an existing paragraph."""
    r = p.add_run(text)
    _set_font(r, name="Consolas", size=10)
    return r


def add_mono_para(doc, text):
    p = doc.add_paragraph()
    add_mono(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_font(r, size=11)
    return p


def add_bullet_rich(doc, parts):
    """Append a bulleted paragraph with a mix of plain + bold + mono runs.
    parts is a sequence of (text, kind) tuples where kind is one of
    'plain', 'bold', 'mono'."""
    p = doc.add_paragraph(style="List Bullet")
    for text, kind in parts:
        r = p.add_run(text)
        if kind == "bold":
            _set_font(r, size=11, bold=True)
        elif kind == "mono":
            _set_font(r, name="Consolas", size=10)
        else:
            _set_font(r, size=11)
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    """Add a Word table with header row + body rows.
    col_widths_cm: optional list of column widths in centimetres."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header row.
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        _set_font(r, size=11, bold=True)

    # Body rows.
    for ri, row in enumerate(rows, start=1):
        for ci, text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            _set_font(r, size=10)

    if col_widths_cm:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)

    return table


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # -------- Page setup ----------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # -------- Title block ---------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Auto Tag User Guide")
    _set_font(r, size=28, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    sub = doc.add_paragraph()
    r = sub.add_run("Kinetic > Audits > Auto Tag")
    _set_font(r, name="Consolas", size=12, color=RGBColor(0x55, 0x55, 0x55))

    ver = doc.add_paragraph()
    r = ver.add_run("Version 6 - May 2026")
    _set_font(r, size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    add_para(
        doc,
        "Auto Tag scans your Revit model for untagged elements across a "
        "configurable set of MEP and equipment categories, places tags "
        "at smart positions next to each element, and writes a QA "
        "report. It is built to be run regularly: a one-button refresh "
        "of all annotation across a discipline rather than a hand-by-"
        "hand tagging pass."
    )

    # -------- What's new in v6 ---------------------------------------------
    doc.add_heading("What's new in v6", level=1)

    add_para(
        doc,
        "Version 6 reshapes how tags get placed and how your settings "
        "are remembered. Highlights:"
    )

    add_bullet_rich(doc, [
        ("Adjacent placement by default.", "bold"),
        (" Tags now sit beside their host element with a leader rather "
         "than on top of it. The position is chosen to avoid clashing "
         "with other tags, text, dimensions, view-crop edges, and key "
         "model elements (walls, framing, columns, ceilings, grids).",
         "plain"),
    ])
    add_bullet_rich(doc, [
        ("Per-project persistence.", "bold"),
        (" Everything you set in the dialog - tag family, length "
         "filters, placement mode, leader preference - is saved into ",
         "plain"),
        ("<project>/.bim/auto_tag.json", "mono"),
        (" when you close the tool. Next launch reads it back, and so "
         "does every other team member who opens the same project.",
         "plain"),
    ])
    add_bullet_rich(doc, [
        ("Three new grid columns - ", "bold"),
        ("Placement", "mono"),
        (", ", "plain"),
        ("Offset mm", "mono"),
        (", ", "plain"),
        ("Side", "mono"),
        (" - exposing the placement controls per row.", "plain"),
    ])
    add_bullet_rich(doc, [
        ("Delete Existing Tags button.", "bold"),
        (" A clean-slate action that wipes all existing tags in the "
         "ticked categories before you re-tag. Confirmed via dialog "
         "and undoable in one Ctrl+Z.", "plain"),
    ])
    add_bullet_rich(doc, [
        ("Leader logic tightened.", "bold"),
        (" When placement is adjacent, the leader is forced on so the "
         "tag visually connects to its host. The ", "plain"),
        ("Lead", "mono"),
        (" checkbox shows ticked and disabled in that mode; switching "
         "back to on-element placement restores whatever you had set.",
         "plain"),
    ])

    # -------- Setup ---------------------------------------------------------
    doc.add_heading("Setup", level=1)

    add_para(doc, "Before first use on a project:")
    add_bullet(
        doc,
        "Ensure the Kinetic extension is installed and registered with "
        "pyRevit. Confirm the Auto Tag button is visible under the "
        "Kinetic ribbon, Audits panel."
    )
    add_bullet(
        doc,
        "Load at least one tag family for every category you intend to "
        "tag. Auto Tag will not place a tag for a category that has no "
        "tag family loaded - it lists them as 'no tag symbol available' "
        "on the report."
    )
    add_bullet(
        doc,
        "If your team uses shared parameters that tags read from (sizes "
        "formatted strings, mounting heights, etc.), confirm they are "
        "loaded into the project before tagging. Auto Tag uses whichever "
        "tag family you select, but the tag family can only display "
        "parameters that exist on the host element."
    )

    # -------- Dialog tour ---------------------------------------------------
    doc.add_heading("The Auto Tag dialog", level=1)

    add_para(
        doc,
        "Click the Auto Tag button to open the dialog. The dialog has "
        "four areas, top to bottom:"
    )

    add_bullet_rich(doc, [
        ("Discipline and subcategories", "bold"),
        (" - pick a discipline on the left, tick one or more "
         "subcategories on the right. Each ticked subcategory becomes "
         "a row in the configuration grid below.", "plain"),
    ])
    add_bullet_rich(doc, [
        ("Tagging profiles grid", "bold"),
        (" - one row per ticked subcategory. Each row carries its own "
         "filters and placement preferences. Cells that do not apply "
         "to a row's category render greyed-out and read-only.",
         "plain"),
    ])
    add_bullet_rich(doc, [
        ("Scope", "bold"),
        (" - 'Scan whole model' toggles between active-view-only and "
         "whole-model scanning. Tags are always placed in the active "
         "view regardless; the toggle only changes which elements are "
         "considered eligible.", "plain"),
    ])
    add_bullet_rich(doc, [
        ("Results pane + action buttons", "bold"),
        (" - shows the scan summary and gives access to Scan, Place, "
         "Delete Existing, Save Report, Save Excel, and Close.",
         "plain"),
    ])

    # -------- Standard workflow ---------------------------------------------
    doc.add_heading("Standard workflow", level=1)

    add_para(
        doc,
        "The typical run looks like this. Steps 4 and 5 are where "
        "v6 differs most from earlier versions."
    )

    workflow = [
        ("1.", "Open the view you want to tag in. Auto Tag places tags "
               "in the active view, so this matters."),
        ("2.", "Click Auto Tag from the Kinetic ribbon."),
        ("3.", "Pick a discipline, tick the subcategories you want to "
               "scan. Each ticked subcategory adds a row to the grid."),
        ("4.", "Tune each row in place. Set min/max length, orientation "
               "flags, size bounds, tag family, placement mode, offset "
               "distance, preferred side, and skip-already-tagged "
               "behaviour. Cells that do not apply to a row's category "
               "are visibly disabled."),
        ("5.", "Decide on scope (active view vs whole model) using the "
               "Scope toggle."),
        ("6.", "Click Scan Model. The results pane shows a per-profile "
               "breakdown - how many were eligible, how many were "
               "excluded by which rule, how many were already tagged."),
        ("7.", "Click Place Tags. Confirm the count when prompted. The "
               "tool wraps placement in a single Revit transaction so "
               "Ctrl+Z undoes the whole batch."),
        ("8.", "Optionally click Save Report (HTML, opens in browser) "
               "or Save Excel (CSV with BOM, opens in Excel) to capture "
               "what was placed and what was skipped."),
        ("9.", "Click Close. Your settings are saved to the project's "
               "auto_tag.json so the next launch and the rest of the "
               "team see the same configuration."),
    ]
    for n, body in workflow:
        p = doc.add_paragraph()
        r = p.add_run(n + " ")
        _set_font(r, size=11, bold=True)
        r = p.add_run(body)
        _set_font(r, size=11)

    # -------- Per-row configuration -----------------------------------------
    doc.add_heading("Per-row configuration", level=1)

    add_para(
        doc,
        "Every ticked subcategory becomes a row in the grid. The columns:"
    )

    columns = [
        ("On",      "Master toggle for the row. Untick to exclude a "
                    "profile from the scan without losing its settings."),
        ("Discipline / Subcategory", "Identity columns - read-only."),
        ("Min Len mm / Max Len mm",  "Length bounds in millimetres for "
                                     "linear MEP. Blank means no bound."),
        ("Horiz / Vert",             "Restrict to runs that are mostly "
                                     "horizontal or mostly vertical. "
                                     "Mutually exclusive."),
        ("Tol mm",                   "Elevation tolerance for the horiz "
                                     "and vert tests. Default 50 mm."),
        ("W / H / Dia min and max",  "Cross-section bounds per category "
                                     "(width and height for ducts, "
                                     "diameter for pipes etc.)."),
        ("Tag family / type",        "Which Revit tag family to use. "
                                     "Default <use Revit default> picks "
                                     "the first loaded tag for the "
                                     "category - usually NOT what you "
                                     "want. Set this explicitly per row "
                                     "to control tag content."),
        ("Placement",                "v6: 'Adjacent' (default - tag "
                                     "beside the host with a leader) "
                                     "or 'On element' (legacy behaviour "
                                     "- tag on top of the host)."),
        ("Offset mm",                "v6: distance from the host when "
                                     "placement is Adjacent. Default 300 "
                                     "mm. Greyed when placement is On "
                                     "element."),
        ("Side",                     "v6: which direction the tag offsets "
                                     "in. Auto picks perpendicular to "
                                     "the run for linear hosts, above-"
                                     "right for point hosts. Explicit "
                                     "values (Above / Below / Left / "
                                     "Right) force a cardinal direction."),
        ("Lead",                     "Show a leader from tag to host. "
                                     "Forced on (and greyed) when "
                                     "placement is Adjacent."),
        ("Skip",                     "When ticked (default), an element "
                                     "that already has a tag in the view "
                                     "is left alone."),
    ]
    add_table(doc, ["Column", "What it does"], columns,
              col_widths_cm=[5.0, 12.0])

    # -------- Placement modes -----------------------------------------------
    doc.add_heading("Placement modes in detail", level=1)

    doc.add_heading("Adjacent (default)", level=2)
    add_para(
        doc,
        "Each tag is placed off-host with a leader. The tool computes "
        "the candidate position based on the row's Offset and Side "
        "settings, then scores it against an in-memory map of nearby "
        "annotation (tags, text, dimensions) and key model elements "
        "(walls, framing, columns, ceilings, grids, levels, detail and "
        "filled regions). If the first candidate clashes, the tool "
        "tries other directions and distances before giving up; if "
        "every candidate clashes, it picks the least-overlapping fall-"
        "back and flags the placement on the report."
    )
    add_para(
        doc,
        "Tags placed earlier in the same run are added to the clash "
        "map immediately, so two adjacent placements never land on top "
        "of each other."
    )

    doc.add_heading("On element (legacy)", level=2)
    add_para(
        doc,
        "Use this when adjacent placement reads as visually "
        "disconnected for a category - for example, very densely "
        "packed point families where 300 mm offsets create more clutter "
        "than the on-top placement would. The tag goes at the host's "
        "origin (curve midpoint for linear runs, location point for "
        "point families) with the leader behaviour from the Lead "
        "checkbox."
    )

    doc.add_heading("Side options", level=2)
    add_para(doc, "The Side dropdown controls which way the tag offsets:")
    add_bullet_rich(doc, [
        ("Auto", "bold"),
        (" - the smart default. For linear hosts (ducts, pipes, cable "
         "trays, conduits) the tag goes perpendicular to the run; for "
         "point hosts the tag goes above-right.", "plain"),
    ])
    add_bullet_rich(doc, [
        ("Above / Below / Left / Right", "bold"),
        (" - the tag offsets in that direction in the view plane, "
         "regardless of host geometry. Useful when you want a "
         "consistent look across a sheet.", "plain"),
    ])

    # -------- Delete Existing Tags ------------------------------------------
    doc.add_heading("Deleting existing tags", level=1)

    add_para(
        doc,
        "The Delete Existing Tags button - red-tinted in the action "
        "row - clears every existing tag in the ticked categories "
        "before you re-tag. Common workflow: a coordination round "
        "changed the model, the prior tagging is now in wrong "
        "positions, you want to wipe and redo cleanly."
    )

    add_para(doc, "Behaviour:")
    add_bullet(
        doc,
        "Operates only on the categories of currently-ticked profiles. "
        "It will not touch tags for categories you have not ticked, "
        "even if they exist in the same view."
    )
    add_bullet(
        doc,
        "Respects the Scope toggle. Unchecked = active view only. "
        "Checked = whole model."
    )
    add_bullet(
        doc,
        "Shows a confirmation dialog listing the affected categories "
        "and scope before running."
    )
    add_bullet(
        doc,
        "Runs in a single Revit transaction. One Ctrl+Z undoes the "
        "whole deletion."
    )
    add_bullet(
        doc,
        "Invalidates any open scan results, so the tool prompts you "
        "to re-scan before placing new tags."
    )

    # -------- Configuration persistence -------------------------------------
    doc.add_heading("Saving and sharing your configuration", level=1)

    p = doc.add_paragraph()
    r = p.add_run("Configuration lives at ")
    _set_font(r, size=11)
    add_mono(p, "<project>/.bim/auto_tag.json")
    r = p.add_run(". When you close the dialog the file is rewritten "
                  "with the current grid state. When you reopen the "
                  "dialog (or another team member opens the project), "
                  "the file is read back and the grid is hydrated with "
                  "the saved values.")
    _set_font(r, size=11)

    add_para(
        doc,
        "If a project has no auto_tag.json yet, Auto Tag falls back to "
        "the extension's default.json - the org-wide baseline shipped "
        "with the Kinetic extension. So a brand-new project still "
        "opens with sensible defaults; saving from any team member "
        "creates the project-specific file."
    )

    add_para(doc, "Shortcuts:", italic=True)
    add_bullet_rich(doc, [
        ("Shift+Click", "bold"),
        (" the Auto Tag button to open ", "plain"),
        ("auto_tag.json", "mono"),
        (" in a text editor without launching the dialog. Useful for "
         "bulk edits or one-off corrections.", "plain"),
    ])

    # -------- Reports and logs ----------------------------------------------
    doc.add_heading("Reports and logs", level=1)

    files = [
        ("<project>/.bim/reports/auto_tag_<timestamp>.html",
         "HTML QA report with one section per profile: filters in "
         "force, eligible / placed counts, per-rule exclusion table, "
         "per-element rows."),
        ("<project>/.bim/reports/auto_tag_<timestamp>.csv",
         "Flat CSV (UTF-8 with BOM, opens cleanly in Excel) with one "
         "row per scanned element. Suitable for pivots and filters."),
        ("<project>/.bim/logs/auto_tag.log",
         "Timestamped log of every scan, placement, deletion, and "
         "config load / save. First place to look when behaviour "
         "surprises you."),
        ("<project>/.bim/auto_tag.json",
         "The persisted configuration. Hand-edit safe; the tool reads "
         "it on next launch."),
    ]
    add_table(doc, ["Path", "Contents"], files, col_widths_cm=[7.5, 9.5])

    add_para(
        doc,
        "When the project is unsaved (no path on disk), all four "
        "locations fall back to your user home folder.",
        italic=True,
    )

    # -------- Troubleshooting -----------------------------------------------
    doc.add_heading("Troubleshooting", level=1)

    qa = [
        ("Tags are still on top of elements after upgrading to v6.",
         "Either pyRevit is running the old cached code (reload the "
         "extension or restart Revit), or the project's auto_tag.json "
         "has a profile with placement_mode set to 'on_element'. "
         "Check the row's Placement column - it should say 'Adjacent'."),
        ("The Placement / Offset / Side columns are not showing.",
         "Same cache issue. Reload pyRevit. If they still do not "
         "appear, scroll horizontally in the grid - on narrow displays "
         "the columns may sit beyond the visible area until you widen "
         "the window."),
        ("Tag content is sparse - shows size only, no system or other "
         "info I expected.",
         "Tag content is defined by the tag family, not Auto Tag. "
         "Check that the Tag family / type column for that row points "
         "at the rich multi-label family rather than the default. "
         "If the rich family is not in the dropdown, it is not "
         "loaded in the project - load it via Insert > Load Family."),
        ("My settings reset every time I open the dialog.",
         "Likely the project has not been saved to disk yet (Revit "
         "can't store .bim/auto_tag.json without a project path), or "
         "the user account does not have write permission to the "
         "project folder. Check the log for save failures."),
        ("Delete Existing Tags removed tags I did not want removed.",
         "Ctrl+Z immediately - the deletion is a single transaction "
         "and is fully reversible. Then check which subcategories "
         "were ticked when you ran the delete; only those categories "
         "are in scope."),
        ("A placement was marked 'fallback' on the report - what does "
         "that mean?",
         "Adjacent placement tried every candidate position and all "
         "of them clashed with something. The tool picked the least-"
         "bad option rather than skipping the tag. Either widen the "
         "offset, change the preferred side, or move the surrounding "
         "annotation manually."),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        r = p.add_run("Q. " + q)
        _set_font(r, size=11, bold=True)
        p = doc.add_paragraph()
        r = p.add_run("A. " + a)
        _set_font(r, size=11)

    # -------- Supported categories ------------------------------------------
    doc.add_heading("Supported categories", level=1)

    cats = [
        ("Cable Trays, Conduits, Ducts, Flex Ducts, Pipes",
         "Linear", "Length, orientation, size, visibility, skip-tagged"),
        ("Pipe Accessories, Mechanical Equipment, Electrical Equipment",
         "Point", "Visibility, skip-tagged, family/system filters"),
        ("Sprinklers, Lighting Fixtures, Plumbing Fixtures",
         "Point", "Visibility, skip-tagged"),
        ("Pipe / Duct / Conduit / Cable Tray Fittings",
         "Point", "Visibility, skip-tagged"),
        ("Duct Accessories, Lighting Devices, Electrical Fixtures, "
         "Fire Alarm Devices, Specialty Equipment",
         "Point", "Visibility, skip-tagged"),
        ("Generic Models", "Point", "Visibility, skip-tagged"),
    ]
    add_table(doc, ["Category", "Geometry", "Rules supported"], cats,
              col_widths_cm=[8.0, 2.5, 6.5])

    # -------- Limitations ---------------------------------------------------
    doc.add_heading("Known limitations", level=1)

    add_bullet(
        doc,
        "Clash detection uses estimated tag bounding boxes (based on "
        "tag family text size and view scale), not measured ones. "
        "Tags with unusually long labels may overflow the estimate "
        "and miss real clashes. Visual review is still recommended on "
        "dense sheets."
    )
    add_bullet(
        doc,
        "Tags can only display parameters that exist on the host "
        "element. If you need information that is not currently a "
        "parameter (mounting height, height from FFL, custom IDs), "
        "you need either (a) a tag family that reads from a different "
        "parameter that does exist, or (b) a separate workflow to "
        "compute and write those values to a shared parameter before "
        "tagging."
    )
    add_bullet(
        doc,
        "Linked-model elements are not tagged. The collector skips "
        "Revit links."
    )
    add_bullet(
        doc,
        "Whole-model scope reports elements eligible elsewhere but "
        "cannot place tags outside the active view; you would need to "
        "open each view and run there."
    )

    # -------- Footer note ---------------------------------------------------
    doc.add_paragraph()
    foot = doc.add_paragraph()
    r = foot.add_run(
        "Kinetic platform - Auto Tag - Annotation QA v6. "
        "Generated from build_user_guide.py; do not hand-edit this "
        "Word file - re-run the generator instead.")
    _set_font(r, size=9, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT)
    print("Wrote", OUTPUT)
    print("Size", os.path.getsize(OUTPUT), "bytes")


if __name__ == "__main__":
    main()
